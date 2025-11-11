#!/usr/bin/env python3
"""
Build Fire Weather Brief - Integration script for Virginia-2025-Fall-Fire-Season
Generates DOCX briefs from live NWS and DOF data
Part of: https://github.com/jamesdcochran-oss/Virginia-2025-Fall-Fire-Season
"""

import json
import sys
import os
from datetime import datetime, timedelta
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def local_points(temp_f=None, rh_min=None, wind_sust=None):
    """
    Calculate local fire danger points based on temperature, RH, and wind.
    Returns total points (0-15).
    """
    # Temperature points
    if temp_f is None:
        t = 0
    elif temp_f > 86:
        t = 5
    elif temp_f >= 76:
        t = 4
    elif temp_f >= 66:
        t = 3
    elif temp_f >= 58:
        t = 2
    elif temp_f >= 50:
        t = 1
    else:
        t = 0
    
    # RH points
    if rh_min is None:
        rh = 0
    elif rh_min <= 19:
        rh = 5
    elif rh_min <= 29:
        rh = 4
    elif rh_min <= 39:
        rh = 3
    elif rh_min <= 59:
        rh = 2
    elif rh_min <= 90:
        rh = 1
    else:
        rh = 0
    
    # Wind points (sustained 20-ft)
    if wind_sust is None:
        w = 0
        print(f"WARNING: Missing wind data, defaulting to 0 points")
    elif wind_sust >= 26:
        w = 5
    elif wind_sust >= 21:
        w = 4
    elif wind_sust >= 16:
        w = 3
    elif wind_sust >= 11:
        w = 2
    else:
        w = 1
    
    return t + rh + w


def class_from_points(pts):
    """Convert points to danger class string."""
    if pts <= 8:
        return "1 (Low)"
    if pts <= 11:
        return "2 (Moderate)"
    if pts <= 15:
        return "3 (High)"
    if pts <= 18:
        return "4 (Very High)"
    return "5 (Extreme)"


def build_doc(data, out_docx):
    """
    Build the Five Forks Fire Weather Brief DOCX from JSON input.
    
    Args:
        data: Dictionary containing fire weather data
        out_docx: Output path for DOCX file
    """
    try:
        # Validate required structure
        if "meta" not in data:
            raise ValueError("Missing 'meta' section in input JSON")
        
        meta = data["meta"]
        
        # Validate dates
        if "dates" not in meta or len(meta["dates"]) < 3:
            raise ValueError("'meta.dates' must contain at least 3 dates")
        
        dates = meta["dates"]
        d1, d2, d3 = dates[0], dates[1], dates[2]
        
        # Parse dates once and reuse (OPTIMIZATION)
        try:
            d1_parsed = datetime.strptime(d1, '%Y-%m-%d')
            d2_parsed = datetime.strptime(d2, '%Y-%m-%d')
            d3_parsed = datetime.strptime(d3, '%Y-%m-%d')
        except ValueError as e:
            raise ValueError(f"Invalid date format in 'meta.dates': {e}")
        
        counties = meta.get("counties", [])
        if not counties:
            print("WARNING: No counties specified in meta.counties")
        
        doc = Document()
        
        # Title
        p = doc.add_paragraph()
        r = p.add_run(meta.get("title", "Five Forks Fire Weather Danger Class Forecast"))
        r.bold = True
        r.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_text = (
            f"Dates: {d1_parsed.strftime('%b %d')}–{d3_parsed.strftime('%d, %Y')}  |  "
            f"Counties: {', '.join(counties)}\n"
            f"Purpose: {meta.get('purpose', 'Fire weather danger assessment')}"
        )
        p2 = doc.add_paragraph(subtitle_text)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # CSI summary
        doc.add_paragraph("🔹 Daily Class Summary")
        csi_table = doc.add_table(rows=1, cols=4)
        csi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["Day", "CSI", "Predicted Class Day", "Points"]
        for i, h in enumerate(headers):
            cell = csi_table.rows[0].cells[i]
            cell.text = h
            try:
                cell.paragraphs[0].runs[0].font.bold = True
            except (IndexError, AttributeError):
                pass
        
        for row in data.get("csi_summary", []):
            r = csi_table.add_row().cells
            r[0].text = row.get("day", "")
            r[1].text = str(row.get("csi", ""))
            r[2].text = row.get("predicted_class_day", "")
            r[3].text = str(row.get("points", ""))
        
        # DOF CSI context
        ctx = data.get("dof_csi_context", {})
        if ctx:
            ctx_text = (
                f"Petersburg (Five Forks District) — "
                f"CSI {ctx.get('petersburg_csi_today', '?')}; "
                f"{ctx.get('narrative', '')}"
            )
            doc.add_paragraph(ctx_text)
        
        doc.add_paragraph()
        
        # County table (7 columns)
        hdrs = [
            "County",
            f"{d1_parsed.strftime('%b %d')} Local",
            f"{d1_parsed.strftime('%b %d')} DOF",
            f"{d2_parsed.strftime('%b %d')} Local",
            f"{d2_parsed.strftime('%b %d')} DOF",
            f"{d3_parsed.strftime('%b %d')} Local",
            f"{d3_parsed.strftime('%b %d')} DOF"
        ]
        
        table = doc.add_table(rows=1, cols=len(hdrs))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, h in enumerate(hdrs):
            cell = table.rows[0].cells[i]
            cell.text = h
            try:
                cell.paragraphs[0].runs[0].font.bold = True
            except (IndexError, AttributeError):
                pass
        
        cinputs = data.get("county_inputs", {})
        
        for county in counties:
            ci = cinputs.get(county, {})
            weather = ci.get("weather", {})
            dof_default = ci.get("dof_class_default", "1 (Low)")
            
            # Helper function for local class calculation
            def local_for(day_str, fallback=None):
                w = weather.get(day_str, {})
                if not w and fallback:
                    w = weather.get(fallback, {})
                    print(f"WARNING: Using fallback data for {county} on {day_str}")
                
                if "local_class" in w:
                    return w["local_class"]
                
                pts = local_points(w.get("temp_max"), w.get("rh_min"), w.get("wind_20ft"))
                return f"{class_from_points(pts)} ({pts} pts)"
            
            d1_local = local_for(d1)
            d2_local = local_for(d2, fallback=d1)
            d3_local = local_for(d3, fallback=d2)
            
            # DOF for each day: allow explicit override per day
            def dof_for(day_str):
                w = weather.get(day_str, {})
                return w.get("dof_class", dof_default)
            
            d1_dof = dof_for(d1)
            d2_dof = dof_for(d2)
            d3_dof = dof_for(d3)
            
            row = table.add_row().cells
            row[0].text = county
            row[1].text = d1_local
            row[2].text = d1_dof
            row[3].text = d2_local
            row[4].text = d2_dof
            row[5].text = d3_local
            row[6].text = d3_dof
        
        doc.add_paragraph()
        
        # Logic recap (3 columns)
        doc.add_paragraph("🔥 Local vs DOF Class Logic Recap")
        t2 = doc.add_table(rows=1, cols=3)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, h in enumerate(["Factor", "Local Class", "DOF Class"]):
            c = t2.rows[0].cells[i]
            c.text = h
            try:
                c.paragraphs[0].runs[0].font.bold = True
            except (IndexError, AttributeError):
                pass
        
        factors = [
            ("Temperature (Max)", "Points 0–5 by threshold; using forecast max", "Included within CSI; smoothed"),
            ("Relative Humidity (Min)", "Points 0–5; sensitive to dips <40%", "Reflected via CSI; less reactive short-term"),
            ("20-ft Wind (sustained)", "Points 0–5; gusts inform Ops notes", "Indirect in CSI"),
            ("Antecedents (DSR, 48h rain, green-up)", "Optional inputs; add for full Steps 1–7", "CSI captures antecedents")
        ]
        
        for f, l, d in factors:
            r = t2.add_row().cells
            r[0].text = f
            r[1].text = l
            r[2].text = d
        
        # Ops notes
        if data.get("ops_notes"):
            doc.add_paragraph()
            doc.add_paragraph("🔍 Operational Notes for the period")
            for note in data["ops_notes"]:
                doc.add_paragraph(f"• {note}")
        
        # Sector blocks
        if data.get("sectors"):
            for key in ["petersburg", "farmville"]:
                sec = data["sectors"].get(key)
                if not sec:
                    continue
                
                doc.add_paragraph()
                doc.add_paragraph(sec.get("label", f"{key.title()} Sector"))
                
                tbl = doc.add_table(rows=1, cols=6)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for i, h in enumerate(["Period", "Sky", "Precip (%) / Type", "Temp/RH", "20-ft Wind", "Notes"]):
                    cell = tbl.rows[0].cells[i]
                    cell.text = h
                    try:
                        cell.paragraphs[0].runs[0].font.bold = True
                    except (IndexError, AttributeError):
                        pass
                
                for p in sec.get("periods", []):
                    row = tbl.add_row().cells
                    row[0].text = p.get("period", "")
                    row[1].text = p.get("sky", "")
                    row[2].text = p.get("precip", "")
                    row[3].text = p.get("temp_rh", "")
                    row[4].text = p.get("wind", "")
                    row[5].text = p.get("notes", "")
        
        doc.add_paragraph()
        foot = doc.add_paragraph(f"Source: {meta.get('source_note', 'NWS + Virginia DOF')}")
        foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        foot.runs[0].font.size = Pt(9)
        
        # Save document
        doc.save(out_docx)
        print(f"✅ Successfully wrote {out_docx}")
        
    except KeyError as e:
        print(f"❌ ERROR: Missing required key in JSON: {e}")
        sys.exit(1)
    except ValueError as e:
        print(f"❌ ERROR: Invalid data: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ ERROR: Unexpected error building document: {e}")
        sys.exit(1)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python scripts/build_five_forks_brief.py <input.json> <output.docx>")
        print("\nExample:")
        print("  python scripts/build_five_forks_brief.py fire_weather.json briefs/$(date +%Y%m%d).docx")
        sys.exit(1)
    
    input_json = sys.argv[1]
    output_docx = sys.argv[2]
    
    # Validate input file exists
    if not os.path.exists(input_json):
        print(f"❌ ERROR: Input file not found: {input_json}")
        sys.exit(1)
    
    try:
        with open(input_json, "r") as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"❌ ERROR: Invalid JSON in {input_json}: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ ERROR: Could not read {input_json}: {e}")
        sys.exit(1)
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_docx)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        print(f"Created directory: {output_dir}")
    
    build_doc(data, output_docx)
